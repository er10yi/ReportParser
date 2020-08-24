# -*- coding: utf-8 -*-
from docx import Document
import xlwt
import sys
import os


# 设置表格样式
def set_style(name, height, bold=False, color=1):
    style = xlwt.XFStyle()
    font = xlwt.Font()
    font.name = name
    font.bold = bold
    font.color_index = 4
    font.height = height

    pattern = xlwt.Pattern()  # 创建一个模式
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN  # 设置其模式为实型
    pattern.pattern_fore_colour = color

    style.pattern = pattern
    borders = xlwt.Borders()
    borders.left = 1
    borders.right = 1
    borders.top = 1
    borders.bottom = 1
    borders.bottom_colour = 0x3A
    style.borders = borders

    style.font = font
    return style


def docx_report_2_excel(docx_name, excel_last_name='漏洞统计'):
    if '.docx' not in docx_name:
        docx_file = docx_name + '.docx'
    else:
        docx_file = docx_name
    if not os.path.isfile(docx_file):
        print(docx_file + ' 不存在')
        exit(1)
    document = Document(docx_file)
    workbook = xlwt.Workbook()
    report_name = docx_name.split('测试报告')[0]
    excel_name = report_name + excel_last_name

    sheet1 = workbook.add_sheet(excel_name, cell_overwrite_ok=True)
    # row0 = ["漏洞分类", "漏洞名称", "风险", "影响范围(位置-链接-参数)", "修复方案", "发现时间", "修复时间", "联系人", "修复人"]
    row0 = ["漏洞分类", "漏洞名称", "风险", "影响范围(位置-链接-参数)", "修复方案", "发现时间", "修复时间"]

    vul_content = []
    for paragraph in document.paragraphs:
        if paragraph.text != '':
            # # print(paralen)
            # # print(paragraph.text)
            vul_content.append(paragraph.text)
    str_content = ",".join(vul_content)
    # # print(str_content)
    split_content = str_content.split('漏洞分类：')
    vuln_name_list = []
    table_map = {}
    for split_str in split_content:
        vul_top = split_str.split(',', 1)[0]
        if len(vul_top) != 1:
            # # print('漏洞分类：' + vul_top)
            vuln_name_list.append(vul_top)
            vul_second = split_str.split(vul_top, 1)[1].split(', 漏洞名称：')
            vul_second_list = []
            for vul_name in vul_second:
                if len(vul_name) != 0 and vul_name != '' and vul_name is not None:
                    # # print('漏洞名称：' + vul_name.replace(',', ''))
                    vul_second_list.append(vul_name.replace(', ', ''))
            table_map.setdefault(vul_top, vul_second_list)
    # # print(split_str + ' ' + str(split_str.count('漏洞名称：')))
    # # print(vuln_name_list)
    # # print(table_map)
    # print('漏洞分类：' + '漏洞名称')
    # for key, value in table_map.items():
    #     print(key + "：", end='')
    #     for name in value:
    #         print(name, end=' ')
    #     print()
    # print('********************')

    critical = 0
    high = 0
    medium = 0
    low = 0
    info = 0

    i = 0
    m = 0
    for j in range(0, len(row0)):
        sheet1.write(0, j, row0[j], set_style('微软雅黑', 230, True, 23))
    for key, value in table_map.items():
        # print('分类：' + key + ' ' + str(len(value)))
        # sheet1.write(i + 1, 0, key, set_style('微软雅黑', 200, True))
        sheet1.write_merge(i + 1, i + len(value), 0, 0, key, set_style('微软雅黑', 200, True))  # 合并行单元格

        # 设置各行宽度
        # 分类
        sheet1.col(0).width = 5000
        # 名称
        sheet1.col(1).width = 6000
        # 风险
        sheet1.col(2).width = 1200
        # 影响范围
        sheet1.col(3).width = 20000
        # 修复方案
        sheet1.col(4).width = 10000
        # 发现时间
        sheet1.col(5).width = 2600
        # 修复时间
        sheet1.col(6).width = 2600

        for name in value:
            # print('名称：' + name)
            sheet1.write(m + 1, 1, name, set_style('微软雅黑', 200))
            str_row_content = []
            for k, row in enumerate(document.tables[i].rows):
                row_content = []
                for cell in row.cells:
                    c = cell.text
                    if c not in row_content:
                        row_content.append(c)
                str_row_content.append(row_content)
                # # print(row_content)
            str_row_content = str(str_row_content).replace("['", "").replace("']", "")
            # # print(str_row_content)
            raw_risk = str_row_content.split('影响范围')[0]
            risk = raw_risk.replace(", ", "").split('风险等级')[1]
            # print('风险：' + risk)
            if risk.strip() == '严重':
                critical = critical + 1
                sheet1.write(m + 1, 2, risk, set_style('微软雅黑', 200, False, 16))
            if risk.strip() == '高危':
                high = high + 1
                sheet1.write(m + 1, 2, risk, set_style('微软雅黑', 200, False, 10))
            if risk.strip() == '中危':
                medium = medium + 1
                sheet1.write(m + 1, 2, risk, set_style('微软雅黑', 200, False, 13))
            if risk.strip() == '低危':
                low = low + 1
                sheet1.write(m + 1, 2, risk, set_style('微软雅黑', 200, False, 15))
            if risk.strip() == '信息':
                info = info + 1
                sheet1.write(m + 1, 2, risk, set_style('微软雅黑', 200, False, 11))

            raw_affect0 = str_row_content.split('影响范围')[1]
            raw_affect = raw_affect0.split('描述及证明')[0]
            array_affect = raw_affect.split(", 位置', '链接', '参数, ")[1].replace("'", " ").split(", ")
            affect_i = 0

            # print('影响范围：')
            affect_cell = []
            # sheet1.write(m + 1, 3, array_affect, set_style('微软雅黑', 200))
            affect_index = 1
            for affect in array_affect:
                if affect and affect_i % 3 == 0:
                    affect_cell.append(str(affect_index) + '.')
                    affect_index += 1
                affect_i = affect_i + 1
                affect_cell.append(affect)
                # print(affect, end=' ')
                if affect_i == 3:
                    # print()
                    affect_cell.append('；')
            # print()
            sheet1.write(m + 1, 3, affect_cell, set_style('微软雅黑', 200, False, 1))

            raw_solution = raw_affect0.split('描述及证明')[1].split('修复方案')[1]
            array_solution = raw_solution.split(", 发现时间, ")[0].replace(", ", "").split('\\n')
            # print('解决方案：')
            # sheet1.write(m + 1, 4, array_solution, set_style('微软雅黑', 200))
            solution_cell = []
            for index, solution in enumerate(array_solution):
                solution_cell.append(str(index + 1) + '.' + solution + '；')
            sheet1.write(m + 1, 4, solution_cell, set_style('微软雅黑', 200))

            # print('发现时间：')
            raw_time = raw_solution.split('发现时间')[1]
            # print(raw_time.split('复测结果')[0].replace(", ", ""))
            sheet1.write(m + 1, 5, raw_time.split('复测结果')[0].replace(", ", ""), set_style('微软雅黑', 200))

            # print('复测结果：')
            raw_retest = raw_time.split('复测结果')[1]
            if '已修复证明' in raw_retest:
                # print('已修复', end=' ')
                # print('修复时间：', end='')
                # print(raw_retest.split('已修复证明')[0].split('\\n')[-1])
                sheet1.write(m + 1, 6, raw_retest.split('已修复证明')[0].split('\\n')[-1].replace(',', ''),
                             set_style('微软雅黑', 200))
            else:
                # print('未修复')
                sheet1.write(m + 1, 6, '', set_style('微软雅黑', 200))
                # sheet1.write(m + 1, 6, '未修复', set_style('微软雅黑', 200, True))

            i += 1
            m += 1

    row00 = ["测试系统", "网站链接", "风险总量", {'严重': critical}, {'高危': high}, {'中危': medium}, {'低危': low}, {'信息': info}]
    total = critical + high + medium + low + info
    sheet2 = workbook.add_sheet(report_name + '漏洞简报', cell_overwrite_ok=True)
    sheet2.col(0).width = 5000
    sheet2.col(1).width = 5000
    sheet2.col(2).width = 2600
    sheet2.col(3).width = 1200
    sheet2.col(4).width = 1200
    sheet2.col(5).width = 1200
    sheet2.col(6).width = 1200
    sheet2.col(7).width = 1200

    # newcol = m + 3
    # newj = 3

    newcol = 0
    newj = 3

    sheet2.write(newcol + 1, 0, report_name, set_style('微软雅黑', 200, False, 1))
    sheet2.write(newcol + 1, 1, '手动填写', set_style('微软雅黑', 200, False, 1))
    sheet2.write(newcol + 1, 2, total, set_style('微软雅黑', 200, False, 1))
    for j in range(0, len(row00)):
        # str
        if type(row00[j]) is str:
            sheet2.write(newcol, j, row00[j], set_style('微软雅黑', 230, False, 1))
        # 字典
        if type(row00[j]) is dict:
            if row00[j].get('严重') is not None and row00[j].get('严重') != 0:
                sheet2.write(newcol, newj, '严重', set_style('微软雅黑', 200, False, 16))
                sheet2.write(newcol + 1, newj, row00[j].get('严重'), set_style('微软雅黑', 200, False, 1))
                newj = newj + 1
            if row00[j].get('高危') is not None and row00[j].get('高危') != 0:
                sheet2.write(newcol, newj, '高危', set_style('微软雅黑', 200, False, 10))
                sheet2.write(newcol + 1, newj, row00[j].get('高危'), set_style('微软雅黑', 200, False, 1))
                newj = newj + 1
            if row00[j].get('中危') is not None and row00[j].get('中危') != 0:
                sheet2.write(newcol, newj, '中危', set_style('微软雅黑', 200, False, 13))
                sheet2.write(newcol + 1, newj, row00[j].get('中危'), set_style('微软雅黑', 200, False, 1))
                newj = newj + 1
            if row00[j].get('低危') is not None and row00[j].get('低危') != 0:
                sheet2.write(newcol, newj, '低危', set_style('微软雅黑', 200, False, 15))
                sheet2.write(newcol + 1, newj, row00[j].get('低危'), set_style('微软雅黑', 200, False, 1))
                newj = newj + 1
            if row00[j].get('信息') is not None and row00[j].get('信息') != 0:
                sheet2.write(newcol, newj, '信息', set_style('微软雅黑', 200, False, 11))
                sheet2.write(newcol + 1, newj, row00[j].get('信息'), set_style('微软雅黑', 200, False, 1))

        # sheet1.write(0, j, row00[j], set_style('微软雅黑', 230, True, 23))

    workbook.save(excel_name + '.xls')
    print(report_name + ' 漏洞简报')
    print('********************')
    print('漏洞分类及名称：')
    # print('漏洞分类：' + '漏洞名称')
    for key, value in table_map.items():
        print(key + "：", end='')
        for name in value:
            print(name, end='；')
        print()
    print('********************')
    print('总数：' + str(total))
    print('严重：' + str(critical))
    print('高危：' + str(high))
    print('中危：' + str(medium))
    print('低危：' + str(low))
    print('信息：' + str(info))


if __name__ == '__main__':
    if len(sys.argv) == 1:
        print('python3 测试报告解析工具')
        print('只支持docx格式')
        print('用法：python report_2_excel.py 报告名字 输出名称\n'
              '例如：python report_2_excel.py 某项目测试报告 漏洞报告\n'
              '或者：python report_2_excel.py 某项目测试报告')
        exit(1)

    if len(sys.argv) == 2:
        docx_report_2_excel(sys.argv[1])

    if len(sys.argv) == 3:
        docx_report_2_excel(sys.argv[1], sys.argv[2])
